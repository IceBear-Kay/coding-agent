import hashlib
import io
import json
import os
import struct
import time
from contextlib import suppress
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

import coding_agent.document_tools as document_tools
from coding_agent.document_tools import (
    MAX_DOCUMENT_PAGES,
    MAX_DOCUMENT_RESULT_BYTES,
    MAX_DOCUMENT_SOURCE_BYTES,
    MAX_DOCUMENT_TEXT_CHARS,
    ReadDocumentArguments,
    read_document_tool_spec,
)
from coding_agent.models import ToolCall, ToolResult
from coding_agent.tools import ToolDispatcher, ToolRegistry, Workspace


def _pdf_bytes(*texts: str) -> bytes:
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    for text in texts:
        page = writer.add_blank_page(width=612, height=792)
        stream = DecodedStreamObject()
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode())
        page[NameObject("/Contents")] = stream
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
        )
    result = io.BytesIO()
    writer.write(result)
    return result.getvalue()


def _dispatch(workspace: Workspace, call: ToolCall) -> ToolResult:
    registry = ToolRegistry([read_document_tool_spec(workspace)])
    return ToolDispatcher(registry).dispatch(call)


def _slow_parse_worker(*_args: object) -> None:
    time.sleep(2)


def _oversized_parse_worker(*args: object) -> None:
    connection = args[-1]
    assert hasattr(connection, "send_bytes")
    connection.send_bytes(b"x" * (MAX_DOCUMENT_RESULT_BYTES + 1))
    connection.close()


def _environment_probe_worker(*args: object) -> None:
    connection = args[-1]
    assert hasattr(connection, "send_bytes")
    document_tools._sanitize_worker_environment()
    payload = json.dumps(
        {
            "text": "probe",
            "marker": os.environ.get("CODING_AGENT_TEST_MODEL_KEY"),
        }
    ).encode("utf-8")
    connection.send_bytes(payload)
    connection.close()


def _partial_parse_worker(*args: object) -> None:
    connection = args[-1]
    assert hasattr(connection, "_send")
    payload = b'{"text":"partial"}'
    connection._send(struct.pack("!i", len(payload)))
    time.sleep(2)
    with suppress(OSError):
        connection._send(payload)


def test_read_document_extracts_pdf_text_and_page_range(tmp_path: Path) -> None:
    target = tmp_path / "notes.pdf"
    target.write_bytes(_pdf_bytes("first page", "second page"))
    digest = hashlib.sha256(target.read_bytes()).hexdigest()

    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(
            id="call_pdf",
            name="read_document",
            arguments={"path": "notes.pdf", "start_page": 2, "end_page": 2},
        ),
    )

    payload = json.loads(result.content)
    assert result.is_error is False
    assert payload["status"] == "completed"
    assert payload["format"] == "pdf"
    assert "second page" in payload["text"]
    assert "first page" not in payload["text"]
    assert payload["processed_start_page"] == 2
    assert payload["processed_end_page"] == 2
    assert payload["total_pages"] == 2
    assert hashlib.sha256(target.read_bytes()).hexdigest() == digest


def test_read_document_pdf_range_can_start_in_middle_and_reach_last_page(
    tmp_path: Path,
) -> None:
    target = tmp_path / "three-pages.pdf"
    target.write_bytes(_pdf_bytes("first", "middle", "last"))

    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(
            id="call_pdf_middle",
            name="read_document",
            arguments={"path": "three-pages.pdf", "start_page": 2, "end_page": 3},
        ),
    )

    payload = json.loads(result.content)
    assert payload["status"] == "completed"
    assert payload["processed_start_page"] == 2
    assert payload["processed_end_page"] == 3
    assert payload["selected_pages"] == 2
    assert "middle" in payload["text"]
    assert "last" in payload["text"]
    assert "first" not in payload["text"]


def test_read_document_pdf_can_read_only_the_last_page(tmp_path: Path) -> None:
    target = tmp_path / "last-page.pdf"
    target.write_bytes(_pdf_bytes("first", "last"))

    payload = json.loads(
        _dispatch(
            Workspace(tmp_path),
            ToolCall(
                id="call_pdf_last",
                name="read_document",
                arguments={"path": target.name, "start_page": 2, "end_page": 2},
            ),
        ).content
    )

    assert payload["status"] == "completed"
    assert payload["processed_start_page"] == 2
    assert payload["processed_end_page"] == 2
    assert payload["selected_pages"] == 1
    assert payload["truncated"] is False
    assert "last" in payload["text"]


def test_read_document_pdf_text_limit_reports_actual_processed_page(tmp_path: Path) -> None:
    target = tmp_path / "text-limit.pdf"
    target.write_bytes(_pdf_bytes("x" * (MAX_DOCUMENT_TEXT_CHARS + 100), "second"))

    payload = json.loads(
        _dispatch(
            Workspace(tmp_path),
            ToolCall(
                id="call_pdf_text_limit",
                name="read_document",
                arguments={"path": target.name},
            ),
        ).content
    )

    assert payload["status"] == "completed"
    assert payload["selected_pages"] == 2
    assert payload["processed_end_page"] == 1
    assert payload["processed_pages"] == 1
    assert payload["truncated"] is True
    assert payload["truncation_reasons"] == ["text_limit"]
    assert "second" not in payload["text"]


def test_read_document_pdf_default_page_limit_is_explicit(tmp_path: Path) -> None:
    target = tmp_path / "many-pages.pdf"
    target.write_bytes(_pdf_bytes(*(f"page-{index}" for index in range(25))))

    payload = json.loads(
        _dispatch(
            Workspace(tmp_path),
            ToolCall(id="call_pdf_default", name="read_document", arguments={"path": target.name}),
        ).content
    )

    assert payload["status"] == "completed"
    assert payload["processed_start_page"] == 1
    assert payload["processed_end_page"] == MAX_DOCUMENT_PAGES
    assert payload["selected_pages"] == MAX_DOCUMENT_PAGES
    assert payload["truncated"] is True
    assert payload["truncation_reasons"] == ["page_limit"]
    assert "page-24" not in payload["text"]


def test_read_document_pdf_no_text_before_unread_page_is_marked_partial(
    tmp_path: Path,
) -> None:
    target = tmp_path / "text-after-limit.pdf"
    target.write_bytes(_pdf_bytes(*([""] * MAX_DOCUMENT_PAGES + ["page-21"])))

    payload = json.loads(
        _dispatch(
            Workspace(tmp_path),
            ToolCall(
                id="call_pdf_no_text_limit", name="read_document", arguments={"path": target.name}
            ),
        ).content
    )

    assert payload["status"] == "no_text"
    assert payload["processed_end_page"] == MAX_DOCUMENT_PAGES
    assert payload["selected_pages"] == MAX_DOCUMENT_PAGES
    assert payload["truncated"] is True
    assert payload["truncation_reasons"] == ["page_limit"]
    assert "unread pages may contain text" in payload["message"]


def test_read_document_pdf_no_text_message_is_limited_to_selected_pages(
    tmp_path: Path,
) -> None:
    target = tmp_path / "partial-text.pdf"
    target.write_bytes(_pdf_bytes("", "text on another page"))

    payload = json.loads(
        _dispatch(
            Workspace(tmp_path),
            ToolCall(
                id="call_pdf_empty_range",
                name="read_document",
                arguments={"path": target.name, "start_page": 1, "end_page": 1},
            ),
        ).content
    )

    assert payload["status"] == "no_text"
    assert payload["selected_pages"] == 1
    assert payload["selected_text_pages"] == 0
    assert "selected PDF pages" in payload["message"]
    assert "unread pages may contain text" in payload["message"]


def test_read_document_extracts_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    target = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    document.add_paragraph("After table")
    document.save(target)

    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_docx", name="read_document", arguments={"path": "notes.docx"}),
    )

    payload = json.loads(result.content)
    assert result.is_error is False
    assert payload["status"] == "completed"
    assert payload["format"] == "docx"
    assert payload["text"] == "Before table\nA\tB\nC\tD\nAfter table"
    assert payload["paragraphs"] == 2
    assert payload["tables"] == 1


@pytest.mark.parametrize(
    ("filename", "content", "status"),
    [
        ("notes.txt", b"plain", "unsupported_format"),
        ("notes.doc", b"old", "unsupported_format"),
        ("notes.docm", b"macro", "unsupported_format"),
        ("broken.pdf", b"not a pdf", "parse_error"),
        ("broken.docx", b"not a zip", "invalid_document"),
    ],
)
def test_read_document_reports_unsupported_and_damaged_documents(
    tmp_path: Path,
    filename: str,
    content: bytes,
    status: str,
) -> None:
    (tmp_path / filename).write_bytes(content)
    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_document", name="read_document", arguments={"path": filename}),
    )

    payload = json.loads(result.content)
    assert result.is_error is True
    assert payload["status"] == status


def test_read_document_reports_empty_and_encrypted_pdf(tmp_path: Path) -> None:
    (tmp_path / "empty.pdf").write_bytes(_pdf_bytes(""))
    empty = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_empty", name="read_document", arguments={"path": "empty.pdf"}),
    )
    assert json.loads(empty.content)["status"] == "no_text"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt("secret")
    encrypted = io.BytesIO()
    writer.write(encrypted)
    (tmp_path / "encrypted.pdf").write_bytes(encrypted.getvalue())
    locked = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_locked", name="read_document", arguments={"path": "encrypted.pdf"}),
    )
    assert json.loads(locked.content)["status"] == "encrypted"


def test_read_document_truncates_text_with_explicit_marker(tmp_path: Path) -> None:
    target = tmp_path / "long.docx"
    document = Document()
    document.add_paragraph("x" * (MAX_DOCUMENT_TEXT_CHARS + 100))
    document.save(target)

    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_long", name="read_document", arguments={"path": "long.docx"}),
    )
    payload = json.loads(result.content)
    assert payload["status"] == "completed"
    assert payload["truncated"] is True
    assert payload["truncation_reasons"] == ["text_limit"]
    assert payload["text"].endswith("...[document text truncated]")
    assert len(payload["text"]) == MAX_DOCUMENT_TEXT_CHARS


def test_read_document_rejects_page_options_for_docx_and_bad_ranges(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("text")
    document.save(tmp_path / "notes.docx")
    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(
            id="call_docx_page",
            name="read_document",
            arguments={"path": "notes.docx", "start_page": 1},
        ),
    )
    assert json.loads(result.content)["status"] == "invalid_arguments"

    with pytest.raises(ValueError, match="before"):
        ReadDocumentArguments(path="notes.pdf", start_page=2, end_page=1)
    with pytest.raises(ValueError, match="maximum"):
        ReadDocumentArguments(
            path="notes.pdf",
            start_page=1,
            end_page=MAX_DOCUMENT_PAGES + 1,
        )
    with pytest.raises(ValueError, match="maximum"):
        ReadDocumentArguments(path="notes.pdf", end_page=MAX_DOCUMENT_PAGES + 1)
    bool_page = _dispatch(
        Workspace(tmp_path),
        ToolCall(
            id="call_bool_page",
            name="read_document",
            arguments={"path": "notes.docx", "start_page": True},
        ),
    )
    assert bool_page.is_error is True
    assert "Invalid arguments for read_document" in bool_page.content


def test_read_document_rejects_path_escape_and_protected_directory(tmp_path: Path) -> None:
    workspace = Workspace(tmp_path, protected_paths=[tmp_path / "sessions"])
    (tmp_path / "sessions").mkdir()
    (tmp_path / "sessions" / "archive.pdf").write_bytes(_pdf_bytes("secret"))

    for path in ("../outside.pdf", "sessions/archive.pdf"):
        result = _dispatch(
            workspace,
            ToolCall(id="call_path", name="read_document", arguments={"path": path}),
        )
        assert result.is_error is True
        assert json.loads(result.content)["status"] == "invalid_path"


def test_read_document_rejects_file_and_parent_symlinks(tmp_path: Path) -> None:
    outside = tmp_path.parent / "outside-document.pdf"
    outside.write_bytes(_pdf_bytes("outside"))
    file_link = tmp_path / "file-link.pdf"
    parent_link = tmp_path / "parent-link"
    try:
        file_link.symlink_to(outside)
        parent_link.symlink_to(outside.parent, target_is_directory=True)
    except OSError:
        pytest.skip("symlink creation is unavailable")

    workspace = Workspace(tmp_path)
    file_result = _dispatch(
        workspace,
        ToolCall(id="call_file_link", name="read_document", arguments={"path": file_link.name}),
    )
    parent_result = _dispatch(
        workspace,
        ToolCall(
            id="call_parent_link",
            name="read_document",
            arguments={"path": "parent-link/outside-document.pdf"},
        ),
    )

    assert file_result.is_error is True
    assert parent_result.is_error is True
    assert json.loads(file_result.content)["status"] == "invalid_path"
    assert json.loads(parent_result.content)["status"] == "invalid_path"


def test_read_document_rejects_source_larger_than_budget(tmp_path: Path) -> None:
    (tmp_path / "large.pdf").write_bytes(b"0" * (MAX_DOCUMENT_SOURCE_BYTES + 1))
    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_large", name="read_document", arguments={"path": "large.pdf"}),
    )
    payload = json.loads(result.content)
    assert payload["status"] == "too_large"


def test_read_document_parser_timeout_terminates_worker(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(document_tools, "_parse_worker", _slow_parse_worker)

    with pytest.raises(document_tools._DocumentReadError, match="time limit"):
        document_tools._parse_document("docx", b"", None, None, timeout_seconds=0.1)

    assert not __import__("multiprocessing").active_children()


def test_read_document_partial_result_respects_overall_deadline(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(document_tools, "_parse_worker", _partial_parse_worker)

    started = time.monotonic()
    with pytest.raises(document_tools._DocumentReadError, match="time limit"):
        document_tools._parse_document("docx", b"", None, None, timeout_seconds=0.2)

    assert time.monotonic() - started < 1.5
    assert not __import__("multiprocessing").active_children()


def test_read_document_rejects_oversized_parser_result(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(document_tools, "_parse_worker", _oversized_parse_worker)

    with pytest.raises(document_tools._DocumentReadError, match="size limit"):
        document_tools._parse_document("docx", b"", None, None, timeout_seconds=2)


def test_read_document_parser_worker_uses_minimal_environment(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("CODING_AGENT_TEST_MODEL_KEY", "fake-marker-not-a-secret")
    monkeypatch.setattr(document_tools, "_parse_worker", _environment_probe_worker)

    result = document_tools._parse_document("docx", b"", None, None, timeout_seconds=2)

    assert result["text"] == "probe"
    assert result["marker"] is None


def test_read_document_rejects_file_replacement_before_open(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    target = tmp_path / "notes.pdf"
    target.write_bytes(_pdf_bytes("inside"))
    outside = tmp_path.parent / "outside-document.pdf"
    outside.write_bytes(_pdf_bytes("outside"))
    original_capture = document_tools._capture_path_state

    def replace_file(path: Path, root: Path) -> object:
        state = original_capture(path, root)
        os.replace(outside, path)
        return state

    monkeypatch.setattr(document_tools, "_capture_path_state", replace_file)
    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(id="call_replaced_file", name="read_document", arguments={"path": target.name}),
    )

    assert result.is_error is True
    assert json.loads(result.content)["status"] == "invalid_path"
    assert json.loads(result.content)["path"] == target.name


def test_read_document_rejects_parent_replacement_before_open(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    parent = tmp_path / "nested"
    parent.mkdir()
    target = parent / "notes.pdf"
    target.write_bytes(_pdf_bytes("inside"))
    replacement = tmp_path / "replacement"
    replacement.mkdir()
    (replacement / "notes.pdf").write_bytes(_pdf_bytes("outside"))
    original_capture = document_tools._capture_path_state

    def replace_parent(path: Path, root: Path) -> object:
        state = original_capture(path, root)
        os.replace(parent, tmp_path / "nested-old")
        os.replace(replacement, parent)
        return state

    monkeypatch.setattr(document_tools, "_capture_path_state", replace_parent)
    result = _dispatch(
        Workspace(tmp_path),
        ToolCall(
            id="call_replaced_parent",
            name="read_document",
            arguments={"path": "nested/notes.pdf"},
        ),
    )

    assert result.is_error is True
    assert json.loads(result.content)["status"] == "invalid_path"
